"""
RAG 文档处理模块

实现文档的解析、清洗、切片等功能。
"""
import re
import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path
import tempfile

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """
    文档切片
    
    存储单个文档切片的内容和元数据。
    """
    content: str
    metadata: Dict[str, Any]
    chunk_id: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "content": self.content,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id
        }


class DocumentCleaner:
    """
    文档清洗器
    
    清洗文档内容，去除无用信息。
    """
    
    # 需要移除的模式
    REMOVE_PATTERNS = [
        r'\x00',  # 空字符
        r'\x0c',  # 换页符
        r'\ufeff',  # BOM 字符
        r'[\r\n]+',  # 多个换行符
        r'[ \t]+',  # 多个空格/制表符
    ]
    
    @classmethod
    def clean(cls, text: str) -> str:
        """
        清洗文本
        
        Args:
            text: 原始文本
        
        Returns:
            清洗后的文本
        """
        # 移除特殊字符
        for pattern in cls.REMOVE_PATTERNS:
            text = re.sub(pattern, ' ', text)
        
        # 移除多余的空白
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # 去除首尾空白
        text = text.strip()
        
        return text
    
    @classmethod
    def remove_headers_footers(cls, text: str) -> str:
        """
        移除页眉页脚
        
        Args:
            text: 原始文本
        
        Returns:
            移除页眉页脚后的文本
        """
        # 简单实现：移除重复的行（可能是页眉页脚）
        lines = text.split('\n')
        
        # 统计每行出现的次数
        line_counts = {}
        for line in lines:
            line = line.strip()
            if line:
                line_counts[line] = line_counts.get(line, 0) + 1
        
        # 移除出现次数过多的行（可能是页眉页脚）
        threshold = max(3, len(lines) // 10)
        cleaned_lines = [
            line for line in lines
            if line.strip() and line_counts.get(line.strip(), 0) < threshold
        ]
        
        return '\n'.join(cleaned_lines)


class TextSplitter:
    """
    文本切片器
    
    将长文本切分成合适大小的片段。
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        separator: str = "\n\n"
    ):
        """
        初始化切片器
        
        Args:
            chunk_size: 切片大小（字符数）
            chunk_overlap: 切片重叠大小（字符数）
            separator: 分隔符
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator
    
    def split(self, text: str) -> List[str]:
        """
        切分文本
        
        Args:
            text: 原始文本
        
        Returns:
            切片列表
        """
        # 先按分隔符切分
        sections = text.split(self.separator)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            section_size = len(section)
            
            # 如果当前切片加上新段落不超过限制，则添加
            if current_size + section_size + len(self.separator) <= self.chunk_size:
                current_chunk.append(section)
                current_size += section_size + len(self.separator)
            else:
                # 保存当前切片
                if current_chunk:
                    chunks.append(self.separator.join(current_chunk))
                
                # 如果段落本身超过限制，需要进一步切分
                if section_size > self.chunk_size:
                    sub_chunks = self._split_large_section(section)
                    chunks.extend(sub_chunks)
                    current_chunk = []
                    current_size = 0
                else:
                    # 开始新切片
                    current_chunk = [section]
                    current_size = section_size
        
        # 保存最后一个切片
        if current_chunk:
            chunks.append(self.separator.join(current_chunk))
        
        # 添加重叠
        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)
        
        return chunks
    
    def _split_large_section(self, section: str) -> List[str]:
        """
        切分过大的段落
        
        Args:
            section: 段落文本
        
        Returns:
            切片列表
        """
        chunks = []
        start = 0
        
        while start < len(section):
            end = start + self.chunk_size
            
            # 尝试在句子边界切分
            if end < len(section):
                # 查找最近的句子边界
                for i in range(end, max(start, end - 100), -1):
                    if section[i] in '。！？.!?.':
                        end = i + 1
                        break
            
            chunk = section[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
        
        return chunks
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """
        添加切片重叠
        
        Args:
            chunks: 原始切片列表
        
        Returns:
            添加重叠后的切片列表
        """
        overlapped_chunks = []
        
        for i, chunk in enumerate(chunks):
            if i > 0:
                # 从上一个切片的末尾提取重叠部分
                prev_chunk = chunks[i - 1]
                overlap_text = prev_chunk[-self.chunk_overlap:]
                
                # 查找句子边界
                for j in range(len(overlap_text) - 1, -1, -1):
                    if overlap_text[j] in '。！？.!?.':
                        overlap_text = overlap_text[j + 1:].strip()
                        break
                
                if overlap_text:
                    chunk = overlap_text + self.separator + chunk
            
            overlapped_chunks.append(chunk)
        
        return overlapped_chunks


class DocumentProcessor:
    """
    文档处理器
    
    完整的文档处理流程：解析 → 清洗 → 切片。
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 切片大小
            chunk_overlap: 切片重叠大小
        """
        self.cleaner = DocumentCleaner()
        self.splitter = TextSplitter(chunk_size, chunk_overlap)
    
    def process(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        处理文档
        
        Args:
            text: 文档文本
            metadata: 文档元数据
        
        Returns:
            文档切片列表
        """
        if metadata is None:
            metadata = {}
        
        # 1. 清洗文档
        cleaned_text = self.cleaner.clean(text)
        cleaned_text = self.cleaner.remove_headers_footers(cleaned_text)
        
        # 2. 切分文档
        chunks_text = self.splitter.split(cleaned_text)
        
        # 3. 创建切片对象
        chunks = []
        for i, chunk_text in enumerate(chunks_text):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_index'] = i
            chunk_metadata['total_chunks'] = len(chunks_text)
            
            chunk = DocumentChunk(
                content=chunk_text,
                metadata=chunk_metadata,
                chunk_id=f"{metadata.get('doc_id', 'unknown')}_{i}"
            )
            chunks.append(chunk)
        
        logger.info(f"Processed document into {len(chunks)} chunks")
        
        return chunks
    
    def process_file(
        self,
        file_path: str,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        处理文件
        
        Args:
            file_path: 文件路径
            text: 文件文本内容
            metadata: 文件元数据
        
        Returns:
            文档切片列表
        """
        if metadata is None:
            metadata = {}
        
        # 添加文件信息
        path = Path(file_path)
        metadata.update({
            'file_path': file_path,
            'file_name': path.name,
            'file_type': path.suffix.lower(),
            'doc_id': path.stem
        })
        
        return self.process(text, metadata)
    
    def process_document(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        处理文档文件
        
        支持多种文档格式：PDF、TXT、DOC、DOCX、XLS、XLSX
        
        Args:
            file_path: 文件路径
            metadata: 文档元数据
        
        Returns:
            文档切片列表
        """
        if metadata is None:
            metadata = {}
        
        path = Path(file_path)
        file_type = path.suffix.lower()
        
        logger.info(f"Processing document: {path.name} (type: {file_type})")
        
        # 根据文件类型选择解析方法
        if file_type == '.txt':
            text = self._parse_txt(file_path)
        elif file_type == '.pdf':
            text = self._parse_pdf(file_path)
        elif file_type in ['.doc', '.docx']:
            text = self._parse_word(file_path)
        elif file_type in ['.xls', '.xlsx']:
            text = self._parse_excel(file_path)
        else:
            # 尝试作为文本文件处理
            try:
                text = self._parse_txt(file_path)
            except Exception as e:
                logger.error(f"Failed to parse file {file_path}: {e}")
                raise ValueError(f"Unsupported file type: {file_type}")
        
        # 处理文本
        return self.process_file(file_path, text, metadata)
    
    def _parse_txt(self, file_path: str) -> str:
        """
        解析文本文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文本内容
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _parse_pdf(self, file_path: str) -> str:
        """
        解析PDF文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文本内容
        """
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
            
            return text
            
        except ImportError:
            logger.warning("PyPDF2 not installed, trying pdfplumber")
            try:
                import pdfplumber
                
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() + "\n\n"
                
                return text
                
            except ImportError:
                logger.error("Neither PyPDF2 nor pdfplumber is installed")
                raise ImportError("Please install PyPDF2 or pdfplumber: pip install PyPDF2 or pip install pdfplumber")
    
    def _parse_word(self, file_path: str) -> str:
        """
        解析Word文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文本内容
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            
            return text
            
        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("Please install python-docx: pip install python-docx")
    
    def _parse_excel(self, file_path: str) -> str:
        """
        解析Excel文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文本内容
        """
        try:
            import pandas as pd
            
            # 读取Excel文件的所有sheet
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                # 读取sheet
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # 添加sheet名称
                text += f"=== Sheet: {sheet_name} ===\n\n"
                
                # 将DataFrame转换为文本
                # 方法1：直接转换为字符串
                text += df.to_string(index=False, na_rep='')
                text += "\n\n"
                
                # 方法2：逐行处理（更详细的格式）
                # for idx, row in df.iterrows():
                #     row_text = " | ".join(str(val) for val in row if pd.notna(val))
                #     if row_text.strip():
                #         text += row_text + "\n"
                # text += "\n"
            
            return text
            
        except ImportError:
            logger.error("pandas and openpyxl not installed")
            raise ImportError("Please install pandas and openpyxl: pip install pandas openpyxl")
        except Exception as e:
            logger.error(f"Failed to parse Excel file {file_path}: {e}")
            raise